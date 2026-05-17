from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading("Phân Tích File: Webpage_2_Mission.py", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("")

# 1. Tổng quan
doc.add_heading("1. Tổng Quan", level=1)
doc.add_paragraph(
    "File Webpage_2_Mission.py là module Python phụ trách render trang 'Mission' của website. "
    "Trang này hiển thị thông tin về các Persona (nhân vật đại diện người dùng) và danh sách thành viên nhóm. "
    "File sử dụng thư viện nội bộ pyhtml để truy vấn cơ sở dữ liệu SQLite và nhúng dữ liệu vào HTML template."
)

# 2. Cấu trúc file
doc.add_heading("2. Cấu Trúc File", level=1)
table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "Thành phần"
hdr[1].text = "Mô tả"
rows = [
    ("import pyhtml", "Import thư viện HTML nội bộ của project"),
    ("DATABASE", 'Hằng số trỏ đến file DB: "Database/persona_team.db"'),
    ("TEMPLATE", 'Hằng số trỏ đến file HTML: "Webpage_2_Mission.html"'),
    ("persona_icon()", "Hàm chuyển nhãn persona → ký tự icon ngắn"),
    ("render_persona_card()", "Hàm tạo HTML cho 1 thẻ persona"),
    ("render_team_rows()", "Hàm tạo HTML cho các hàng bảng thành viên"),
    ("get_page_html()", "Hàm chính: đọc template, truy vấn DB, ghép HTML"),
]
for r in rows:
    row = table.add_row().cells
    row[0].text = r[0]
    row[1].text = r[1]

doc.add_paragraph("")

# 3. Chi tiết từng hàm
doc.add_heading("3. Chi Tiết Từng Hàm", level=1)

doc.add_heading("3.1 persona_icon(label)", level=2)
doc.add_paragraph(
    "Nhận vào nhãn persona (ví dụ: 'PERSONA 1') và trả về ký tự viết tắt tương ứng "
    "('P1', 'P2', 'H1', 'H2'). Nếu không khớp thì trả về 'U' (Unknown). "
    "Dùng để hiển thị avatar chữ trong thẻ persona."
)

doc.add_heading("3.2 render_persona_card(persona)", level=2)
doc.add_paragraph(
    "Nhận vào 1 tuple dữ liệu persona gồm 7 trường: name, role, age, location, description, "
    "persona_label, border_color. Trả về chuỗi HTML dạng <article> chứa avatar, tên, thông tin meta, "
    "mô tả và tag nhãn. Màu viền và tag được truyền động qua class CSS (border-{color}, tag-{color})."
)

doc.add_heading("3.3 render_team_rows(team_members)", level=2)
doc.add_paragraph(
    "Nhận vào danh sách các thành viên nhóm, mỗi phần tử là tuple gồm: name, student_id, sub_task, pages. "
    "Trả về chuỗi HTML gồm nhiều thẻ <tr> để nhúng vào bảng team trong trang HTML."
)

doc.add_heading("3.4 get_page_html(form_data)", level=2)
doc.add_paragraph(
    "Đây là hàm chính được gọi từ bên ngoài. Quy trình hoạt động:"
)
steps = [
    "Đọc nội dung file HTML template (Webpage_2_Mission.html)",
    "Truy vấn bảng 'personas' từ SQLite để lấy 4 persona, sắp xếp theo id",
    "Truy vấn bảng 'team_members' từ SQLite để lấy danh sách thành viên",
    "Chia 4 persona thành 2 nhóm: group1 (2 đầu) và group2 (2 còn lại)",
    "Thay thế placeholder {{GROUP1}}, {{GROUP2}}, {{TEAM_MEMBERS}} trong template bằng HTML thực",
    "Trả về chuỗi HTML hoàn chỉnh để render ra trình duyệt",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"  {i}. {s}")

doc.add_paragraph("")

# 4. Luồng dữ liệu
doc.add_heading("4. Luồng Dữ Liệu", level=1)
doc.add_paragraph("SQLite DB  →  pyhtml.get_results_from_query()  →  render functions  →  HTML template  →  Trang web")

doc.add_paragraph("")

# 5. SQL Queries
doc.add_heading("5. Câu Truy Vấn SQL", level=1)

doc.add_heading("Query 1 – Personas:", level=2)
doc.add_paragraph("SELECT name, role, age, location, description, persona_label, border_color\nFROM personas\nORDER BY id;")

doc.add_heading("Query 2 – Team Members:", level=2)
doc.add_paragraph("SELECT name, student_id, sub_task, pages\nFROM team_members\nORDER BY id;")

doc.add_paragraph("")

# 6. Nhận xét
doc.add_heading("6. Nhận Xét", level=1)
notes = [
    ("Điểm mạnh", "Tách biệt rõ logic render và dữ liệu. Dùng template HTML thay vì hardcode. Không dùng JS — đúng yêu cầu project."),
    ("Cần lưu ý", "Hàm get_page_html() nhận form_data nhưng không sử dụng — có thể bỏ tham số này nếu không cần về sau."),
    ("Bảo mật", "Câu SQL dùng chuỗi tĩnh, không có user input nên không có rủi ro SQL injection."),
]
for label, content in notes:
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(content)

doc.save("Phan_Tich_Webpage_2_Mission.docx")
print("Done! File saved: Phan_Tich_Webpage_2_Mission.docx")
